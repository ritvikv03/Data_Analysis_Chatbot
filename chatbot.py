# run 'pip install streamlit google-generativeai pdfplumber pandas python-docx openpyxl numpy' in terminal to install dependencies
# run streamlit run chatbot.py in terminal to start the app

# chatbot.py
# Run:
# pip install streamlit google-generativeai pdfplumber pandas python-docx openpyxl numpy reportlab
# streamlit run chatbot.py

import streamlit as st
import google.generativeai as genai
import pdfplumber
import pandas as pd
import numpy as np
from docx import Document
from io import BytesIO
from datetime import datetime
import os

# --- Optional Dependencies (V1 Feature) ---
try:
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas as pdf_canvas
    PDF_AVAILABLE = True
except Exception:
    PDF_AVAILABLE = False

# -------------------------
# Safe rerun helper (V1 Feature)
# -------------------------
def safe_rerun():
    """Safe rerun for Streamlit."""
    try:
        if hasattr(st, "rerun"):
            st.rerun()
    except Exception:
        pass

# -------------------------
# Page config & Theme CSS (V1 Feature)
# -------------------------
st.set_page_config(page_title="Data Analytics Helper", layout="wide", initial_sidebar_state="expanded")

BG = "#0b1020"
PANEL = "#0f1724"
TEXT = "#e6edf3"
MUTED = "#97a0b3"
ACCENT = "#10a37f"

st.markdown(f"""
<style>
[data-testid="stAppViewContainer"] {{ background-color: {BG}; color: {TEXT}; }}
[data-testid="stSidebar"] {{ background-color: {PANEL}; color: {TEXT}; }}
h1,h2,h3,h4,h5,p,label,div {{ color: {TEXT} !important; }}
a {{ color: {ACCENT} !important; }}
.upload-box {{ background-color: #0f1620; border-radius: 12px; padding: 12px; border: 1px solid #17212a; }}
.file-entry {{ padding: 8px; border-radius: 8px; margin-bottom: 8px; background-color: #09101a; border: 1px solid #16202a; }}
.chat-window {{ background-color: #071021; border-radius: 12px; padding: 16px; height: 520px; overflow-y: auto; }}
.msg-bubble {{ padding: 10px 14px; border-radius: 12px; display: inline-block; max-width:80%; margin-bottom:10px; }}
.msg-user {{ background: rgba(255,255,255,0.06); border: 1px solid rgba(255,255,255,0.04); color: {TEXT}; float: right; clear: both; }}
.msg-bot {{ background: linear-gradient(90deg, rgba(16,163,127,0.12), rgba(16,163,127,0.06)); border: 1px solid rgba(16,163,127,0.18); color: {TEXT}; float: left; clear: both; }}
.clearfix::after {{ content: ""; clear: both; display: table; }}
.input-area {{ background-color: #0c1118; padding: 10px; border-radius: 12px; border: 1px solid #202631; }}
.stButton>button {{ background-color: {ACCENT}; color: #022016; font-weight: 600; border-radius: 8px; padding: 8px 12px; border: none; }}
.stButton>button:hover {{ filter: brightness(0.95); }}
.kb-hint {{ color: {MUTED}; font-size:12px; }}
body {{ background-color: {BG}; color: {TEXT}; }}
.stSidebar {{ background-color: {PANEL}; color: {TEXT}; }}
</style>
""", unsafe_allow_html=True)

# -------------------------
# Session State Init (V1 Feature, enhanced for V2 data)
# -------------------------
if "gemini_api_key" not in st.session_state:
    st.session_state["gemini_api_key"] = ""
if "uploads" not in st.session_state:
    st.session_state["uploads"] = []
if "messages" not in st.session_state:
    st.session_state["messages"] = []
if "model_name" not in st.session_state:
    st.session_state["model_name"] = "gemini-2.5-flash"
if "model_temperature" not in st.session_state:
    st.session_state["model_temperature"] = 0.7
# New state for storing the full context string for the chat model
if "current_file_context" not in st.session_state:
    st.session_state["current_file_context"] = ""
if "current_file_name" not in st.session_state:
    st.session_state["current_file_name"] = ""
if "current_df" not in st.session_state:
    st.session_state["current_df"] = None
if "current_stats" not in st.session_state:
    st.session_state["current_stats"] = None


# -------------------------
# Utility functions (V1 and V2 combined)
# -------------------------
def set_gemini_api_key(api_key: str):
    st.session_state["gemini_api_key"] = api_key
    try:
        genai.configure(api_key=api_key)
    except Exception:
        pass

def add_message(role: str, content: str):
    ts = datetime.utcnow().isoformat()
    # Check for duplicate timestamps before adding (Streamlit rerun safety)
    if not st.session_state["messages"] or st.session_state["messages"][-1].get("ts") != ts:
        st.session_state["messages"].append({"role": role, "content": content, "ts": ts})


# --- V2: Statistical Analysis Functions ---
def analyze_dataset_statistics(df):
    """Perform comprehensive statistical analysis on dataset (V2 Logic)"""
    analysis = {
        'basic_stats': {},
        'distributions': {},
        'correlations': None,
    }
    
    analysis['basic_stats'] = {
        'shape': df.shape,
        'missing_pct': (df.isnull().sum() / len(df) * 100).to_dict(),
        'dtypes': df.dtypes.to_dict(),
        'memory_usage': df.memory_usage(deep=True).sum() / 1024**2
    }
    
    numerical_cols = df.select_dtypes(include=[np.number]).columns.tolist()
    if numerical_cols:
        analysis['distributions'] = df[numerical_cols].describe().to_dict()
        if len(numerical_cols) > 1:
            analysis['correlations'] = df[numerical_cols].corr()
    
    categorical_cols = df.select_dtypes(include=['object']).columns.tolist()
    analysis['categorical'] = {
        col: df[col].value_counts().head(10).to_dict() 
        for col in categorical_cols[:5]
    }
    
    return analysis, numerical_cols, categorical_cols

def extract_text_from_file(uploaded_file):
    """Extract text and perform initial data analysis (V2 Logic)"""
    file_content = ""
    file_type = ""
    stats_analysis = None
    df = None

    try:
        # Use BytesIO to handle the uploaded file object content (V1/V2 compatibility)
        file_stream = BytesIO(uploaded_file.read())
        file_stream.seek(0)
        
        # Determine file type
        name = uploaded_file.name
        mime = uploaded_file.type

        if mime == "text/plain" or name.endswith(".txt"):
            file_content = file_stream.read().decode("utf-8", errors="ignore")
            file_type = "txt"
            
        elif mime == "application/pdf" or name.endswith(".pdf"):
            with pdfplumber.open(file_stream) as pdf:
                file_content = '\n\n'.join([
                    f"Page {i+1}:\n{page.extract_text()}" 
                    for i, page in enumerate(pdf.pages) 
                    if page.extract_text()
                ])
            file_type = "pdf"
            
        elif mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" or name.endswith(".docx"):
            doc = Document(file_stream)
            file_content = '\n\n'.join([
                paragraph.text for paragraph in doc.paragraphs 
                if paragraph.text.strip()
            ])
            file_type = "docx"
            
        elif name.endswith(('.csv', '.xlsx', '.xls')):
            # Read into DataFrame
            if name.endswith('.csv'):
                df = pd.read_csv(file_stream)
                file_type = "csv"
            elif name.endswith(('.xlsx', '.xls')):
                df = pd.read_excel(file_stream)
                file_type = "xlsx"
            
            # Generate statistics
            stats_analysis, num_cols, cat_cols = analyze_dataset_statistics(df)
            
            # Build the text content for the model
            file_content = f"""Dataset Overview:
- Shape: {df.shape[0]} rows × {df.shape[1]} columns
- Memory Usage: {stats_analysis['basic_stats']['memory_usage']:.2f} MB
- Numerical Features: {len(num_cols)} ({', '.join(num_cols[:10])})
- Categorical Features: {len(cat_cols)} ({', '.join(cat_cols[:10])})

Missing Values Analysis (Percent):
{pd.Series(stats_analysis['basic_stats']['missing_pct']).to_string()}

Statistical Summary (Numerical Features):
{df[num_cols].describe().to_string() if num_cols else 'No numerical features'}

Correlation Analysis (Top Features):
{stats_analysis['correlations'].to_string() if stats_analysis['correlations'] is not None else 'Not applicable'}

Sample Data (First 5 Rows):
{df.head(5).to_string()}"""
            
        return file_content, file_type, df, stats_analysis
        
    except Exception as e:
        st.error(f"Error reading file: {e}")
        return None, None, None, None

def generate_ml_stats_analysis(content, file_type, stats_analysis=None):
    """Generate ML/Stats-focused comprehensive analysis using Gemini (V2 Logic)"""
    
    # Configuration is loaded from session state (V1 Setting)
    model = genai.GenerativeModel(
        st.session_state["model_name"],
        generation_config={
            'temperature': st.session_state["model_temperature"],
            'max_output_tokens': 4096,
        }
    )
    
    # Customize prompt based on file type
    if file_type in ['csv', 'xlsx'] and stats_analysis:
        prompt_type = "Dataset"
        prompt_sections = """
1. **EXECUTIVE SUMMARY** (150 words): Overview, potential uses, key statistical insights.
2. **STATISTICAL ANALYSIS**: Distribution patterns, feature correlations (highlight top 3-5), data quality, and preprocessing needs.
3. **MACHINE LEARNING INSIGHTS**: Best ML algorithms (top 3), feature engineering suggestions, potential target variables, and expected challenges.
4. **PRACTICE PROBLEMS** (10 questions): Mixed statistical, ML, and evaluation problems.
5. **QUICK IMPLEMENTATION GUIDE**: Python code starter template, key libraries, and common pitfalls."""
    else:
        prompt_type = "Academic Material (Lecture Notes/Paper)"
        prompt_sections = """
1. **EXECUTIVE SUMMARY** (150 words): Main topics, key concepts, and real-world applications.
2. **CONCEPT BREAKDOWN** (Top 3-5 concepts): Intuitive explanation, key formulas, and common misconceptions for each.
3. **MATHEMATICAL ESSENTIALS**: Core equations (in LaTeX) with explanations, key assumptions, and computational complexity.
4. **PRACTICE QUESTIONS** (12 questions): Conceptual, mathematical, and application-based problems.
5. **IMPLEMENTATION TIPS**: Python code patterns, best practices, and debugging advice."""

    prompt = f"""You are a world-class Machine Learning and Statistics professor. Analyze this {prompt_type} and provide expert guidance.

Content to analyze (limited to first 20k characters for context window):
{content[:20000]}

Provide a focused analysis structured exactly with the following sections (be concise but insightful):
{prompt_sections}"""

    try:
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        st.error(f"Error generating analysis: {e}")
        return None


# --- V1: Export Functions ---
def export_chat_as_docx(chat_history):
    doc = Document()
    doc.add_heading("Chat Export - Data Analytics Helper", level=1)
    for msg in chat_history:
        p = doc.add_paragraph()
        ts_clean = datetime.fromisoformat(msg['ts']).strftime("%Y-%m-%d %H:%M:%S") if 'ts' in msg else "N/A"
        p.add_run(f"{msg['role'].capitalize()} ({ts_clean}): ").bold = True
        p.add_run(msg['content'])
    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return out

def export_chat_as_pdf(chat_history):
    buffer = BytesIO()
    c = pdf_canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    y = height - 72
    c.setFont("Helvetica-Bold", 14)
    c.drawString(72, y, "Chat Export - Data Analytics Helper")
    y -= 28
    c.setFont("Helvetica", 10)
    for msg in chat_history:
        ts_clean = datetime.fromisoformat(msg['ts']).strftime("%Y-%m-%d %H:%M:%S") if 'ts' in msg else "N/A"
        text_lines = f"{msg['role'].capitalize()} ({ts_clean}): {msg['content']}".split('\n')
        for line in text_lines:
            if y < 72:
                c.showPage()
                y = height - 72
                c.setFont("Helvetica", 10)
            c.drawString(72, y, line)
            y -= 14
    c.save()
    buffer.seek(0)
    return buffer


# -------------------------
# API Key Input (V1 Structure)
# -------------------------
if not st.session_state["gemini_api_key"]:
    st.sidebar.success("💰 100% FREE - Powered by Google Gemini")
    st.sidebar.info("📊 1,500 requests/day available")
    st.header("Set up your Gemini API Key")
    api_key_col1, api_key_col2 = st.columns([4,1])
    with api_key_col1:
        api_key = st.text_input("Enter your Gemini API Key:", type="password", key="api_input")
    with api_key_col2:
        st.markdown("<br>", unsafe_allow_html=True)
        if st.button("Set API Key", type="primary"):
            if api_key:
                set_gemini_api_key(api_key)
                st.success("API Key set successfully!")
                safe_rerun()
            else:
                st.error("Please enter a valid API key.")
    with st.expander("How to get your FREE Gemini API Key"):
        st.markdown("""
            1. Go to Google AI Studio (https://aistudio.google.com/app/apikey)
            2. Sign in with Google
            3. Create API Key and paste it above
            """)
    st.stop()
else:
    try:
        genai.configure(api_key=st.session_state["gemini_api_key"])
    except Exception:
        pass

# -------------------------
# Sidebar + Navigation (V1 Structure)
# -------------------------
with st.sidebar:
    st.title("Navigation Hub")
    nav = st.radio("Navigate to", ["Chat", "Uploads", "Analysis", "Settings"], index=0) # Changed Reference to Analysis
    st.divider()
    
    st.header("Current File")
    if not st.session_state["current_file_name"]:
        st.markdown("<div class='file-entry'>No file loaded</div>", unsafe_allow_html=True)
    else:
        name = st.session_state["current_file_name"]
        st.markdown(f"<div class='file-entry'><strong>{name}</strong><br><span style='color:{MUTED}; font-size:12px;'>Context Loaded</span></div>", unsafe_allow_html=True)

    st.divider()
    st.markdown("### Quick Actions")
    if st.button("Clear loaded file & context"):
        st.session_state["current_file_context"] = ""
        st.session_state["current_file_name"] = ""
        st.session_state["current_df"] = None
        st.session_state["current_stats"] = None
        st.session_state["messages"] = []
        st.success("Context cleared.")
        safe_rerun()
    if st.button("Clear chat history"):
        st.session_state["messages"] = []
        st.success("Chat history cleared.")
        safe_rerun()
    st.divider()
    st.markdown("<div class='kb-hint'>Tip: For best results, clear the loaded file before uploading a new one.</div>", unsafe_allow_html=True)

# -------------------------
# Main Pages
# -------------------------
st.title("Data Analytics Chatbot 🤖")
st.caption("Upload lecture notes, datasets, or statistical analysis documents, then ask targeted questions.")

# NAV Pages
if nav == "Uploads":
    st.header("Upload Materials")
    st.markdown("Drop a file to upload. Supported: txt, pdf, docx, csv, xlsx")
    
    # Check if a file is already loaded
    if st.session_state["current_file_name"]:
        st.warning(f"File **{st.session_state['current_file_name']}** is already loaded. Clear it from the sidebar before uploading a new one.")
    else:
        uploaded_file = st.file_uploader(
            "Choose file to analyze",
            accept_multiple_files=False, 
            type=['txt','pdf','docx','csv','xlsx']
        )
        
        if uploaded_file:
            with st.spinner("🔍 Performing advanced analysis..."):
                file_content, file_type, df, stats_analysis = extract_text_from_file(uploaded_file)
            
            if file_content:
                # Store the results in session state
                st.session_state["current_file_context"] = file_content
                st.session_state["current_file_name"] = uploaded_file.name
                st.session_state["current_df"] = df
                st.session_state["current_stats"] = stats_analysis
                st.session_state["messages"] = [] # Clear chat on new upload
                
                st.success(f"✅ File **{uploaded_file.name}** loaded and context extracted. Go to the **Analysis** tab to view the expert report or the **Chat** tab to ask questions.")
                
                # Optional: Show a quick summary after upload
                if df is None:
                    st.text_area("Document Preview (First 500 characters):", file_content[:500], height=150)
                
                # Rerun to update sidebar
                safe_rerun()
            else:
                 st.error("❌ Could not extract content from the file.")


elif nav == "Analysis":
    st.header("Expert ML/Stats Analysis")
    
    if not st.session_state["current_file_context"]:
        st.info("Upload a file in the **Uploads** tab to generate an automated expert analysis report.")
    else:
        st.subheader(f"Report for: {st.session_state['current_file_name']}")
        
        # Check if the detailed analysis has been run before (optional caching)
        if "generated_analysis" not in st.session_state or st.session_state.generated_analysis_file != st.session_state["current_file_name"]:
            
            with st.spinner("⚡ Generating comprehensive expert report..."):
                analysis_result = generate_ml_stats_analysis(
                    st.session_state["current_file_context"], 
                    "csv" if st.session_state["current_df"] is not None else "doc", # Simple type for prompt logic
                    st.session_state["current_stats"]
                )
            
            if analysis_result:
                st.session_state["generated_analysis"] = analysis_result
                st.session_state["generated_analysis_file"] = st.session_state["current_file_name"]
                
        if "generated_analysis" in st.session_state:
            st.markdown(st.session_state["generated_analysis"])
            
            st.divider()
            col1, col2 = st.columns(2)
            with col1:
                st.download_button(
                    label="📥 Download Study Guide (Markdown)",
                    data=st.session_state["generated_analysis"],
                    file_name=f"{st.session_state['current_file_name']}_expert_analysis.md",
                    mime="text/markdown"
                )
            if st.session_state["current_df"] is not None:
                with col2:
                     st.download_button(
                        label="⬇️ Download Data Preview (.csv)",
                        data=st.session_state["current_df"].head(1000).to_csv().encode('utf-8'),
                        file_name=f"{st.session_state['current_file_name']}_preview.csv",
                        mime="text/csv"
                    )

elif nav == "Settings":
    st.header("Settings")
    
    # Model Configuration (V1 Setting)
    st.subheader("Model Configuration")
    st.session_state["model_name"] = st.selectbox(
        "Select Gemini Model",
        ["gemini-2.5-flash", "gemini-2.5-pro"],
        index=["gemini-2.5-flash", "gemini-2.5-pro"].index(st.session_state["model_name"]),
        help="Flash is faster and cheaper, Pro is more capable for complex reasoning."
    )
    st.session_state["model_temperature"] = st.slider(
        "Creativity (Temperature)",
        min_value=0.0, max_value=1.0, value=st.session_state["model_temperature"], step=0.05,
        help="Lower values (closer to 0) for factual, deterministic responses. Higher values for more creative, diverse responses."
    )
    st.info("Changing these settings will not affect the current analysis report, only future chat questions.")

    st.subheader("Chat Export Options")
    col1, col2 = st.columns(2)
    with col1:
        docx_file = export_chat_as_docx(st.session_state["messages"])
        st.download_button("Export Chat to DOCX", docx_file, file_name="chat_export.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    with col2:
        if PDF_AVAILABLE:
            pdf_file = export_chat_as_pdf(st.session_state["messages"])
            st.download_button("Export Chat to PDF", pdf_file, file_name="chat_export.pdf", mime="application/pdf")
        else:
            st.info("PDF export unavailable (reportlab missing)")

    st.subheader("Data Reset")
    if st.button("Reset all session data (API Key will remain)", on_click=lambda: [st.session_state.pop(key, None) for key in list(st.session_state.keys()) if key not in ["gemini_api_key"]]):
        st.success("All non-API-Key session data cleared.")
        safe_rerun()


else:
    # CHAT PAGE (V1 UI + V2 Chat Logic)
    st.header("Chat with the Data Analytics Bot")
    
    # Display chat messages (V1 Styling)
    chat_box = st.container() 
    with chat_box:
        # Show suggested questions if context is loaded but chat is empty
        if not st.session_state.messages and st.session_state["current_file_context"]:
            st.info("💡 **Suggested questions to get started:**\n- What are the main concepts in this file?\n- Can you explain [concept] in simpler terms?\n- What practice problems can you give me?\n- How would I implement this in Python?")
        elif not st.session_state.messages:
            st.info("💡 **Try asking general ML/Stats questions** or upload a file in the **Uploads** tab.")
            
        for msg in st.session_state["messages"]:
            role = msg['role']
            content = msg['content']
            if role == "user":
                st.markdown(f"<div class='msg-bubble msg-user clearfix'>{content}</div>", unsafe_allow_html=True)
            else:
                st.markdown(f"<div class='msg-bubble msg-bot clearfix'>{content}</div>", unsafe_allow_html=True)
    
    # Input box and Send button
    st.markdown("<br><br>", unsafe_allow_html=True)
    with st.container():
        user_input = st.text_area("Your message:", "", key="chat_input", height=80)
        
        if st.button("Send", key="send_button", type="primary"):
            user_prompt = user_input.strip()
            if user_prompt:
                # 1. Add user message to state
                add_message("user", user_prompt)
                
                # 2. Get bot response
                with st.spinner("Gemini is analyzing..."):
                    
                    # Core Gemini Logic for Chat
                    model = genai.GenerativeModel(
                        st.session_state["model_name"],
                        generation_config={'temperature': st.session_state["model_temperature"], 'max_output_tokens': 2048}
                    )
                    
                    # System Context (V2 Logic)
                    context_message = """You are a world-class Machine Learning and Statistics expert with PhD-level knowledge. Provide clear, accurate, and insightful answers.

When answering:
1. Be concise and thorough but simple
2. Use examples and analogies
3. Include relevant formulas when helpful in LaTeX format
4. Give practical implementation advice
5. If referring to the uploaded file, state the name of the file."""
            
                    # File Context (V2 Logic)
                    file_context = ""
                    if st.session_state["current_file_context"]:
                        name = st.session_state["current_file_name"]
                        file_context = f"\n\nContext from loaded file ({name}):\n{st.session_state['current_file_context'][:3000]}"
                    
                    # Build Conversation History
                    conversation_history = ""
                    for msg in st.session_state.messages:
                        conversation_history += f"\n{msg['role']}: {msg['content']}"

                    # Full Prompt
                    full_prompt = f"{context_message}{file_context}\n\nConversation:\n{conversation_history}\n\nassistant:"
                    
                    try:
                        response = model.generate_content(full_prompt)
                        bot_response = response.text
                    except Exception as e:
                        bot_response = f"An error occurred: {e}. Check your API Key or try a simpler question."

                # 3. Add bot message to state
                add_message("bot", bot_response)
                
                # 4. Clear input and rerun
                safe_rerun()
            else:
                st.warning("Please enter a message.")

    st.markdown("---")
    st.markdown("<div style='color:#98a0b6; font-size:12px;'>Tip: Upload files and ask specific questions. This demo is powered by Google Gemini.</div>", unsafe_allow_html=True)
